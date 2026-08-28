import io
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional
from pydantic import BaseModel


class ParsedPage(BaseModel):
    page_number: int
    text: str
    section_title: Optional[str] = None


class ParsedDocument(BaseModel):
    pages: List[ParsedPage]
    full_text: str
    page_count: int
    metadata: Dict[str, Any] = {}


class DocumentParser(ABC):
    @abstractmethod
    def parse(self, file_content: bytes, filename: str) -> ParsedDocument:
        pass


class PDFParser(DocumentParser):
    def parse(self, file_content: bytes, filename: str) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf is required for PDF parsing.")

        reader = PdfReader(io.BytesIO(file_content))
        pages: List[ParsedPage] = []
        full_text_list = []

        total_pages = len(reader.pages)
        for idx, page in enumerate(reader.pages):
            page_num = idx + 1
            text = page.extract_text() or ""
            text = text.strip()
            pages.append(ParsedPage(page_number=page_num, text=text))
            if text:
                full_text_list.append(text)

        full_text = "\n\n".join(full_text_list)
        return ParsedDocument(
            pages=pages,
            full_text=full_text,
            page_count=total_pages,
            metadata={"filename": filename, "format": "pdf"}
        )


class DOCXParser(DocumentParser):
    def parse(self, file_content: bytes, filename: str) -> ParsedDocument:
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing.")

        doc = docx.Document(io.BytesIO(file_content))
        paragraphs_text = []
        current_section = None
        pages: List[ParsedPage] = []

        # DOCX doesn't have fixed hard page breaks, so we treat sections/paragraphs cleanly
        current_page_text = []
        page_counter = 1

        for p in doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue
            if p.style.name.startswith("Heading"):
                current_section = txt
            current_page_text.append(txt)
            paragraphs_text.append(txt)

            # Approximate logical page grouping every ~500 words
            if len(" ".join(current_page_text).split()) >= 450:
                pages.append(
                    ParsedPage(
                        page_number=page_counter,
                        text="\n".join(current_page_text),
                        section_title=current_section
                    )
                )
                page_counter += 1
                current_page_text = []

        if current_page_text:
            pages.append(
                ParsedPage(
                    page_number=page_counter,
                    text="\n".join(current_page_text),
                    section_title=current_section
                )
            )

        if not pages:
            pages.append(ParsedPage(page_number=1, text=""))

        full_text = "\n\n".join(paragraphs_text)
        return ParsedDocument(
            pages=pages,
            full_text=full_text,
            page_count=len(pages),
            metadata={"filename": filename, "format": "docx"}
        )


class TXTParser(DocumentParser):
    def parse(self, file_content: bytes, filename: str) -> ParsedDocument:
        try:
            text = file_content.decode("utf-8")
        except UnicodeDecodeError:
            text = file_content.decode("latin-1", errors="replace")

        # Split into logical ~500 word pages
        paragraphs = text.split("\n\n")
        pages: List[ParsedPage] = []
        current_page_paras = []
        page_counter = 1

        for para in paragraphs:
            para_clean = para.strip()
            if not para_clean:
                continue
            current_page_paras.append(para_clean)
            if len(" ".join(current_page_paras).split()) >= 450:
                pages.append(
                    ParsedPage(
                        page_number=page_counter,
                        text="\n\n".join(current_page_paras)
                    )
                )
                page_counter += 1
                current_page_paras = []

        if current_page_paras:
            pages.append(
                ParsedPage(
                    page_number=page_counter,
                    text="\n\n".join(current_page_paras)
                )
            )

        if not pages:
            pages.append(ParsedPage(page_number=1, text=text.strip()))

        return ParsedDocument(
            pages=pages,
            full_text=text,
            page_count=len(pages),
            metadata={"filename": filename, "format": "txt"}
        )


def get_document_parser(filename: str, mime_type: Optional[str] = None) -> DocumentParser:
    lower_fn = filename.lower()
    if lower_fn.endswith(".pdf") or mime_type == "application/pdf":
        return PDFParser()
    elif lower_fn.endswith(".docx") or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return DOCXParser()
    elif lower_fn.endswith(".txt") or lower_fn.endswith(".md") or (mime_type and "text/" in mime_type):
        return TXTParser()
    else:
        raise ValueError(f"Unsupported file format for {filename}. Supported formats: PDF, DOCX, TXT, MD.")
